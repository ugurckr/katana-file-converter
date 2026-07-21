"""Metin dönüştürücüleri: docx -> txt/md, html -> pdf/md, txt/md -> docx, txt -> pdf.

python-docx, html2text, markdown-it-py ve xhtml2pdf ile çalışır; harici
bir araç gerekmez.
"""

import html as html_escape
from pathlib import Path

import html2text
from docx import Document
from markdown_it import MarkdownIt
from xhtml2pdf import pisa

from .base import register

_md = MarkdownIt("commonmark", {"html": True})

# Word'ün yerleşik başlık stillerini Markdown '#' seviyelerine eşler.
_HEADING_LEVELS = {f"Heading {i}": i for i in range(1, 7)} | {"Title": 1}


def _html_to_pdf(html: str, src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    with dst.open("wb") as f:
        result = pisa.CreatePDF(html, dest=f)
    if result.err:
        raise RuntimeError(f"'{src}' PDF'e dönüştürülürken hata oluştu.")


@register(".docx", ".txt", "Düz metin (metin çıkarma)")
def docx_to_txt(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    lines = [para.text for para in doc.paragraphs]
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


@register(".docx", ".md", "Markdown belge")
def docx_to_md(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        level = _HEADING_LEVELS.get(style)
        if level:
            lines.append(f"{'#' * level} {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n\n".join(lines) + "\n", encoding="utf-8")


@register(".html", ".pdf", "PDF belge")
def html_to_pdf(src: Path, dst: Path) -> None:
    _html_to_pdf(src.read_text(encoding="utf-8"), src, dst)


@register(".html", ".md", "Markdown belge")
def html_to_md(src: Path, dst: Path) -> None:
    converter = html2text.HTML2Text()
    converter.body_width = 0  # satırları yapay olarak kırma
    markdown = converter.handle(src.read_text(encoding="utf-8"))
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(markdown.strip() + "\n", encoding="utf-8")


@register(".txt", ".pdf", "PDF belge")
def txt_to_pdf(src: Path, dst: Path) -> None:
    body = html_escape.escape(src.read_text(encoding="utf-8"))
    html = (
        "<!doctype html>\n<meta charset=\"utf-8\">\n"
        f"<pre style=\"font-family: Helvetica; white-space: pre-wrap;\">{body}</pre>\n"
    )
    _html_to_pdf(html, src, dst)


@register(".txt", ".docx", "Word belgesi")
def txt_to_docx(src: Path, dst: Path) -> None:
    doc = Document()
    for line in src.read_text(encoding="utf-8").splitlines():
        doc.add_paragraph(line)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


@register(".md", ".docx", "Word belgesi")
def md_to_docx(src: Path, dst: Path) -> None:
    """Başlık, paragraf, liste ve kod bloklarını Word'e çevirir;
    kalın/italik düz metin olarak aktarılır."""
    tokens = _md.parse(src.read_text(encoding="utf-8"))
    doc = Document()

    list_style = None
    for i, token in enumerate(tokens):
        if token.type == "bullet_list_open":
            list_style = "List Bullet"
        elif token.type == "ordered_list_open":
            list_style = "List Number"
        elif token.type in ("bullet_list_close", "ordered_list_close"):
            list_style = None
        elif token.type == "heading_open":
            level = min(int(token.tag[1]), 6)
            doc.add_heading(tokens[i + 1].content, level=level)
        elif token.type == "inline" and tokens[i - 1].type == "paragraph_open":
            if list_style:
                doc.add_paragraph(token.content, style=list_style)
            else:
                doc.add_paragraph(token.content)
        elif token.type in ("fence", "code_block"):
            para = doc.add_paragraph()
            run = para.add_run(token.content.rstrip("\n"))
            run.font.name = "Courier New"

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
